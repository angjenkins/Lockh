from bottle import Bottle, run, template,request,TEMPLATE_PATH,static_file,HeaderDict,BaseResponse,response,redirect, error, request

import time
import urllib
import ast
import sys

import LogFeedback
import FileServiceFeedback
import LogActivity
import requests
import os
import P3Auth.LogActivity
import P3Auth.pycas
import P3Search
import StringIO
import xlsxwriter
import re
from uuid import uuid4
 
#from ppGuiConfig import URLToPPSearchApiSolr,GoogleAnalyticsInclusionScript,\
#     LocalURLToRecordFeedback,CAS_SERVER,CAS_PROXY,CAS_RETURN_SERVICE_URL,CAS_LEVEL_OF_ASSURANCE,CAS_LEVEL_OF_ASSURANCE_PREDICATE, TokenTimeout,\
#     OMBMAX_REDIRECT_URL


#import masterconfig
from configs.commonconfigs import CAS_SERVER,CAS_PROXY,CAS_LEVEL_OF_ASSURANCE,CAS_LEVEL_OF_ASSURANCE_PREDICATE, TokenTimeout, autherror,error_msg
from configs.ppGuiConfig import GoogleAnalyticsInclusionScript,LocalURLToRecordFeedback
import socket

machinename = socket.gethostname()
print 'MachineName = ', machinename

if machinename == 'fcoh1c-pp01d':
   from configs.dev.nodeconfig import URLToPPSearchApiSolr, OMBMAX_REDIRECT_URL,CAS_RETURN_SERVICE_URL,CAS_CREATE_SESSION_IF_AUTHENTICATED,FileUploadServiceURL,AuthServiceURL
#test-web-01 & test-web-02
elif machinename in ['fcoh1c-ppweb01t','fcoh1c-ppweb02t']:
   from configs.test.nodeconfig import  URLToPPSearchApiSolr, OMBMAX_REDIRECT_URL,CAS_RETURN_SERVICE_URL,CAS_CREATE_SESSION_IF_AUTHENTICATED,FileUploadServiceURL,AuthServiceURL
#prod-web-01 & prod-web-02
elif machinename in ['fcoh1c-ppweb01p','fcoh1c-ppweb01p']:
   from configs.prod.nodeconfig import  URLToPPSearchApiSolr, OMBMAX_REDIRECT_URL,CAS_RETURN_SERVICE_URL,CAS_CREATE_SESSION_IF_AUTHENTICATED,FileUploadServiceURL,AuthServiceURL

import P3Auth.auth
import cPickle as pickle
from cStringIO import StringIO

# I am duplicating this because I don't really know how t organize python 
# classes.  Probably it should be removed.
import morris_config

import os
import cgi
import time
import urllib
import urlparse

URL_TO_MORRIS_PORTFOLIOS_API = "http://localhost:" + str(morris_config.BOTTLE_DEORATOR_PORTFOLIOS_API_PORT)

URL_TO_MORRIS_TAGS_API = "http://localhost:" + str(morris_config.BOTTLE_DEORATOR_TAGS_API_PORT)


PathToBottleWebApp = "./"
PathToExternalFiles = "../"


PathToCSSFiles=PathToExternalFiles+"css/"

#app = Bottle()
from beaker.middleware import SessionMiddleware
session_opts = {
	'session.type':'memory',
	#'session.cookie_expires':30,
	'session.cookie_expires':TokenTimeout,
	'session.auto':True,
        'session.secure':True
}

app = SessionMiddleware(Bottle(),session_opts)

PricesPaidAPIUsername=None
PricesPaidAPIPassword=None
PricesPaidAPIBasicAuthUsername=None
PricesPaidAPIBasicAuthPassword=None
P3APISALT = None
PYCAS_SECRET = None

def readCredentials():
    global PricesPaidAPIUsername
    if (PricesPaidAPIUsername is None):
        global PricesPaidAPIPassword
        global PricesPaidAPIBasicAuthUsername
        global PricesPaidAPIBasicAuthPassword
        global P3APISALT
        global PYCAS_SECRET
        PricesPaidAPIUsername=os.environ.get("PricesPaidAPIUsername")
        PricesPaidAPIPassword=os.environ.get("PricesPaidAPIPassword")
        PricesPaidAPIBasicAuthUsername=os.environ.get("PricesPaidAPIBasicAuthUsername")
        PricesPaidAPIBasicAuthPassword=os.environ.get("PricesPaidAPIBasicAuthPassword")
        P3APISALT=os.environ.get("P3APISALT")
        PYCAS_SECRET=os.environ.get("PYCAS_SECRET")

# Begin Common Template Strings
FEEDBACK_HTML = template('Feedback')
FOOTER_HTML = template('Footer')
COLUMN_DROPDOWN_HTML = template('ColumnDropdown')
EXTRA_LOGIN_METHODS = template('ExtraLoginMethods')
PORTFOLIO_PANEL = template('PortfolioPanel')
MAINJS_INCLUDES = template('MainJSIncludes')
SLICKGRID_INCLUDES = template('SlickGridIncludes')
JQPLOT_INCLUDES = template('JQPlotIncludes')

# End Common Template Strings

@app.wrap_app.route('/theme/<path:path>')
def server_static(path):
    return static_file(path, root=PathToBottleWebApp+"theme/")

@app.wrap_app.route('/css/<filename>')
def server_static(filename):
    return static_file(filename, root=PathToCSSFiles)

@app.wrap_app.route('/js/<filename>')
def server_static(filename):
    return static_file(filename, root=PathToBottleWebApp+"js/")

@app.wrap_app.route('/MorrisDataDecorator/js/<filename>')
def server_static(filename):
    return static_file(filename, root="../MorrisDataDecorator/js/")
@app.wrap_app.route('/MorrisDataDecorator/imgs/<filename>')
def server_static(filename):
    return static_file(filename, root="../MorrisDataDecorator/imgs/")
@app.wrap_app.route('/MorrisDataDecorator/css/<filename>')
def server_static(filename):
    return static_file(filename, root="../MorrisDataDecorator/css/")


from bottle import template

#@app.wrap_app.route('/')
def legalNotice():
    P3Auth.LogActivity.logPageTurn("nosession","LegalNotice",request.remote_addr)
    return template('LegalNotice',goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/SearchHelp')
def searchHelp():
    P3Auth.LogActivity.logPageTurn("nosession","SearchHelp",request.remote_addr)
    return template('SearchHelp',goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/files',method='GET')
def pptriv():
    return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/files',method='POST')
def pptriv():
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')
    filename = 'REVERSE_AUCTIONS_FY2014_FYQ1.xlsx'
    filename = request.forms.get('filename').upper()
    filename = filename.strip()
    if filename[0:4] == 'AT&T':
	filename = 'AT_T_' + filename[5:]  
    filename_s = filename.split('.')
    filename = filename_s[0] + '.' + filename_s[1].lower()  
    return static_file(filename,'/web/www/docs/files',download=filename)


@app.wrap_app.route('/DataSets',method='GET')
def pptriv():
    return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/DataSets',method='POST')
def datasets():
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf,request.remote_addr)):
        return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)
    
    P3Auth.auth.update_acsrf(ses_id,request.remote_addr)

    P3Auth.LogActivity.logPageTurn(ses_id,"DatasetPage",request.remote_addr)
    return template('datasets',\
                    acsrf=P3Auth.auth.get_acsrf(ses_id),\
                    session_id=ses_id,footer_html=FOOTER_HTML,feedback_html=FEEDBACK_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/DataLoads',method='GET')
def datasets():
    return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/DataLoads',method='POST')
def datasets():
    print 'Calling dataload page'
    username = request.forms.get('username')
    password = request.forms.get('password')

    FileServiceFeedback.logServiceCalled('File Service Called', username, request.remote_addr)

    import re
    if re.match('^[^0-9a-zA-Z]+$',username):
        return template('Login',message='DATALOAD-The username and password you entered did not match our records. Please double-check and try again.',
		    ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

    r = requests.post(AuthServiceURL,auth=(username,password))

    if r.content == 'VALIDATED':
	s = request.environ.get('beaker.session')		
	token = str(uuid4())
	s[token] = {username:password}
	s.save()
	ses_id = P3Auth.auth.create_session_id()
    else:
        return template('Login',message='DATALOAD-The username and password you entered did not match our records. Please double-check and try again.',
		    ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)
    
    return template('dataload',
		    RESTCALL_SUCCESS='U',
                    DISPLAY_MESSAGE='',
		    session_id=ses_id,
		    token=token,
		    acsrf=P3Auth.auth.get_acsrf(ses_id),
		    footer_html=FOOTER_HTML,feedback_html=FEEDBACK_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/FileUploadService',method='GET')
def datasets():
    return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/FileUploadService',method='POST')
def datasets():

    token = request.forms.get('token')
    acsrf = request.forms.get('acsrf')
    ses_id = request.forms.get('session_id')

    uploadfile = request.files.get('File Upload')

    tmp_file_upload = '/web/tmp_file_upload/'

    #We are using beaker to store uname/pwd, the key is the token
    s = request.environ.get('beaker.session')		

    if token in s:
        print 'token = ', s[token]
	username = s[token].keys()[0]
	password = s[token].values()[0]
    else:
	print 'token expired'
        return template('Login',message='DATALOAD-Your session has been inactive for 20 minutes and you have been logged out. Please re-enter your Username and Password to continue.',
		    ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

    name, ext = os.path.splitext(uploadfile.raw_filename)

    FileServiceFeedback.logFileName(uploadfile.raw_filename,request.remote_addr)

    try:
     	with open(tmp_file_upload + uploadfile.raw_filename,'w') as open_file:
     		open_file.write(uploadfile.file.read())
    except Exception,err:
    	import traceback
        #traceback.print_exception(*exc_info)
        traceback.print_exc()

    from requests_toolbelt import MultipartEncoder
    import gzip

    if (bytesto(os.path.getsize(tmp_file_upload + uploadfile.raw_filename),'m') > 150):
	restcall_success="N"
	display_message="Filesize is more than 150MB. Please contact p3@gsa.gov for support."
	return template('dataload',\
                    session_id=request.forms.get('acsrf'),
                    token=token,
                    acsrf=request.forms.get('session_id'),
                    RESTCALL_SUCCESS=restcall_success,
                    DISPLAY_MESSAGE=display_message,
                    footer_html=FOOTER_HTML,feedback_html=FEEDBACK_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

  

    f_in = open(tmp_file_upload + uploadfile.raw_filename,'rb')
    f_out = gzip.open(tmp_file_upload + uploadfile.raw_filename+'.gz','wb')
    f_out.writelines(f_in)
    f_out.close()
    f_in.close()

    file = uploadfile.raw_filename+'.gz'

    FileServiceFeedback.logGeneric("Calling the REST WebService",request.remote_addr)
    m =        MultipartEncoder(fields= { 'filename' :(file ,open(tmp_file_upload + file,"rb") ) })
    #r = requests.post('https://extcws.fas.gsa.gov/soatest/RESTFileAttachmentService/RESTFileIngestAPI/UploadServlet',
    # 		data=m,auth=(username,password),headers={'Content-Type':m.content_type}a

    FileServiceFeedback.logServiceURL(FileUploadServiceURL,request.remote_addr)

    r = requests.post(FileUploadServiceURL,data=m,auth=(username,password),headers={'Content-Type':m.content_type})
    
    FileServiceFeedback.logResponseHeaders(r,request.remote_addr)

    os.remove(tmp_file_upload +uploadfile.raw_filename)
    os.remove(tmp_file_upload +uploadfile.raw_filename+'.gz')

    restcall_success="N"
    display_message = ""
 
    FileServiceFeedback.logGeneric("Calling the REST WebService - Complete",request.remote_addr)


    if r.status_code == 200:
	    if r.content.strip() == "File successfully uploaded":
		restcall_success="Y"
	    else:
		restcall_success="N"
		if r.content.strip() == "Authentication error: Authentication failed: the supplied credential is not right":
			display_message = autherror
		else:
			display_message=r.content.strip()
    elif r.status_code == 500:
              try:
                  display_message = error_msg[r.content.strip()]
	          display_message = 'We received an error trying to upload your file. Please contact p3@gsa.gov for support.'
              except:
		   display_message = 'We received an error trying to upload your file. Please contact p3@gsa.gov for support.'
    elif r.status_code == 401:
              display_message = error_msg[r.content.strip()]
	      display_message = 'We received an error trying to upload your file. Please contact p3@gsa.gov for support.'
    else:
	      display_message = 'We received an error trying to upload your file. Please contact p3@gsa.gov for support.'

    return template('dataload',\
		    session_id=request.forms.get('acsrf'),
		    token=token,
		    acsrf=request.forms.get('session_id'),
		    RESTCALL_SUCCESS=restcall_success,
                    DISPLAY_MESSAGE=display_message,
		    footer_html=FOOTER_HTML,feedback_html=FEEDBACK_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)


@app.wrap_app.route('/CheckSession',method='POST')
def check_valid_session():
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')
    valid_session = "N"
    if P3Auth.auth.is_valid_acsrf(ses_id,acsrf,request.remote_addr):
       valid_session = "Y"
    return valid_session


@app.wrap_app.route('/Logout',method='POST')
def logoutViaPost():
    P3Auth.LogActivity.logPageTurn("nosession","Logout",request.remote_addr)

    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')
    P3Auth.auth.del_session(ses_id)
    return template('Logout',goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/Logout',method='GET')
def logoutViaGet():
    P3Auth.LogActivity.logPageTurn("nosession","Logout",request.remote_addr)

    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')
    P3Auth.auth.del_session(ses_id)
    return template('Logout',goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/')
def login():
    print 'remote addr =', request.remote_addr
    P3Auth.LogActivity.logPageTurn("nosession","LoginPage",request.remote_addr)
    return template('Login',message='F',
	            ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    footer_html=FOOTER_HTML,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/LoginViaMax')
def loginViaMax():
    print 'LoginViaMAx request.method = ', request.method
    P3Auth.LogActivity.logPageTurn("nosession","MaxLoginPage",request.remote_addr)
    response.status = 303 
    domain,path = urlparse.urlparse(CAS_RETURN_SERVICE_URL)[1:3]
    secure=1
    setCookieCommand = P3Auth.pycas.make_pycas_cookie("gateway",domain,path,secure)
    strip = setCookieCommand[12:]
    response.set_header('Set-Cookie', strip)
    opt=""
    location = P3Auth.pycas.get_url_redirect_as_string(CAS_SERVER,CAS_RETURN_SERVICE_URL,opt,secure)
    response.set_header('Location',location)
    return "You will be redirected."+strip+location

@app.wrap_app.route('/ReturnLoginViaMax')
def returnLoginViaMax():
    print 'ReturnLoginViaMAx request.method = ', request.method
    P3Auth.LogActivity.logPageTurn("nosession","ReturnMaxLoginPage",request.remote_addr)

    PYCAS_SECRET=os.environ.get("PYCAS_SECRET")

    P3Auth.LogActivity.logDebugInfo("PYCAS_SECRET:"+PYCAS_SECRET,request.remote_addr)
    ticket = request.query['ticket']
    P3Auth.LogActivity.logDebugInfo("MAX AUTHENTICATED ticket :"+ticket,request.remote_addr)
    status, id, cookie = P3Auth.pycas.check_authenticated_p(CAS_LEVEL_OF_ASSURANCE_PREDICATE,ticket,CAS_SERVER,CAS_PROXY, 
        PYCAS_SECRET, CAS_RETURN_SERVICE_URL, lifetime=None, secure=1, protocol=2, path="/", opt="")
    maxAuthenticatedProperly = (status == P3Auth.pycas.CAS_OK);

    P3Auth.LogActivity.logDebugInfo("MAX AUTHENTICATED WITH ID:"+id,request.remote_addr)

    username = "billybob"
    if (maxAuthenticatedProperly):
        return doStartPageAuthenticated(username)
    else:
        P3Auth.LogActivity.logBadCredentials(username+":failed to Authenticate with Max",request.remote_addr)
# It would be better to make this message configuration in the same way that CAS_LEVEL_OF_ASSURANCE_PREDICATE is...
# But that is for another day.
        if ticket == "":
		return template('Login',message='F',
				 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
			    footer_html=FOOTER_HTML,
			    extra_login_methods=EXTRA_LOGIN_METHODS,
				goog_anal_script=GoogleAnalyticsInclusionScript)
	else:
		return template('Login',message='T',
				 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
			    footer_html=FOOTER_HTML,
			    extra_login_methods=EXTRA_LOGIN_METHODS,
				goog_anal_script=GoogleAnalyticsInclusionScript)

#@app.wrap_app.route('/ReturnLoginViaMax',method='GET')
#def pptriv():
#    print 'request.status =', dir(request.headers)
#    print 'response.status =', response.status
#    return template('Login',message='F',
#			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
#                    extra_login_methods=EXTRA_LOGIN_METHODS,
#                    footer_html=FOOTER_HTML,
#                    goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/StartPage',method='GET')
def pptriv():
    return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

    
@app.wrap_app.route('/StartPage',method='POST')
def pptriv():
    username = request.forms.get('username')
    password = request.forms.get('password')
    # just a little throttle to slow down any denial of service attack..
    time.sleep(1.0);

    readCredentials()


    if (not P3Auth.auth.does_authenticate(username,password,P3APISALT,request.remote_addr)):
        P3Auth.LogActivity.logBadCredentials(username,request.remote_addr)
        return template('Login',message='Improper Credentials.',
		 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    footer_html=FOOTER_HTML,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                        goog_anal_script=GoogleAnalyticsInclusionScript)
    return doStartPageAuthenticated(username)

def doStartPageAuthenticated(username):
    search_string = request.forms.get('search_string')
    search_string = search_string if search_string is not None else ""
    psc_pattern = request.forms.get('psc_pattern')
    ses_id = P3Auth.auth.create_session_id()
    P3Auth.LogActivity.logSessionBegin(username,ses_id,request.remote_addr)
    P3Auth.LogActivity.logPageTurn(ses_id,"StartPage",request.remote_addr)
    return template('StartPage',search_string=search_string,\
                    acsrf=P3Auth.auth.get_acsrf(ses_id),\
                    username=username, \
                    session_id=ses_id,\
                    footer_html=FOOTER_HTML,\
                    psc_pattern=psc_pattern,goog_anal_script=GoogleAnalyticsInclusionScript,timeout=TokenTimeout)

@app.wrap_app.route('/StartPageReturned',method='GET')
def StartPageReturned():
    return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/StartPageReturned',method='POST')
def StartPageReturned():
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf,request.remote_addr)):
        return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

    search_string = request.forms.get('search_string')
    search_string = search_string if search_string is not None else ""
    psc_pattern = request.forms.get('psc_pattern')
    ses_id = P3Auth.auth.create_session_id()
    P3Auth.LogActivity.logPageTurn(ses_id,"StartPageReturned",request.remote_addr)
    return template('StartPage',search_string=search_string,\
                    acsrf=P3Auth.auth.get_acsrf(ses_id),\
                    session_id=ses_id,\
                    footer_html=FOOTER_HTML,\
                    psc_pattern=psc_pattern,goog_anal_script=GoogleAnalyticsInclusionScript,timeout=TokenTimeout)

@app.wrap_app.route('/PricesPaid',method='GET')
def pptriv(): 
    return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/PricesPaid',method='POST')
def pptriv():
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')
    return render_main_page(acsrf,ses_id,request)

def render_main_page(acsrf,ses_id,request):
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf,request.remote_addr)):
        return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
goog_anal_script=GoogleAnalyticsInclusionScript)
    
    P3Auth.auth.update_acsrf(ses_id,request.remote_addr)

    search_string = request.forms.get('search_string')
    search_string = search_string if search_string is not None else ""
    commodity_id = request.forms.get('commodity_id')

    P3Auth.LogActivity.logPageTurn(ses_id,"MainPage",request.remote_addr)
    return template('MainPage',search_string=search_string,\
                    acsrf=P3Auth.auth.get_acsrf(ses_id),\
                    session_id=ses_id,\
                    feedback_url=LocalURLToRecordFeedback,\
                    footer_html=FOOTER_HTML,\
                    feedback_html=FEEDBACK_HTML,\
                    portfolio_panel=PORTFOLIO_PANEL,\
                    column_dropdown=COLUMN_DROPDOWN_HTML,\
                    mainjs_includes=MAINJS_INCLUDES,\
                    slickgrid_includes=SLICKGRID_INCLUDES,\
                    jqplot_includes=JQPLOT_INCLUDES,\
                    commodity_id=commodity_id,goog_anal_script=GoogleAnalyticsInclusionScript)

@app.wrap_app.route('/PortfolioPage',method='POST')
def render_portfolio():
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        return template('Login',message='Improper Credentials or Timeout.',
                    extra_login_methods=EXTRA_LOGIN_METHODS,
                    footer_html=FOOTER_HTML,
                    goog_anal_script=GoogleAnalyticsInclusionScript)

    P3Auth.auth.update_acsrf(ses_id,request.remote_addr)

    P3Auth.LogActivity.logPageTurn(ses_id,"Portfolio",request.remote_addr)

    portfolio = request.forms.get('portfolio')

    return template('Portfolio',acsrf=P3Auth.auth.get_acsrf(ses_id),\
                    session_id=ses_id,\
                    portfolio=portfolio,\
                    feedback_url=LocalURLToRecordFeedback,\
                    footer_html=FOOTER_HTML,\
                    portfolio_panel=PORTFOLIO_PANEL,\
                    column_dropdown=COLUMN_DROPDOWN_HTML,\
                    mainjs_includes=MAINJS_INCLUDES,\
                    slickgrid_includes=SLICKGRID_INCLUDES,\
                    jqplot_includes=JQPLOT_INCLUDES,\
                        goog_anal_script=GoogleAnalyticsInclusionScript)


@app.wrap_app.route('/returnPortfolio',method='POST')
def apisolr():
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')

    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        dict = {0: {"status": "BadAuthentication"}}
        return dict;

    portfolio = request.forms.get('portfolio')

    print "portfolio = "+portfolio
    r = requests.get(URL_TO_MORRIS_PORTFOLIOS_API+"/decoration/"+portfolio)
    content = r.text
    d = ast.literal_eval(r.text)
    p3ids = d['data']

    readCredentials()

    payload = { 'username' : PricesPaidAPIUsername,\
                                'password' : PricesPaidAPIPassword,\
                                'p3ids' : pickle.dumps(p3ids)
                }

    readCredentials()

    r = requests.post(URLToPPSearchApiSolr+"/fromIds", data=payload, \
                          auth=(PricesPaidAPIBasicAuthUsername, PricesPaidAPIBasicAuthPassword), verify=False)

    P3Auth.LogActivity.logDebugInfo("Got Past Post to :"+URLToPPSearchApiSolr,request.remote_addr)

    content = r.text

    # This is inefficient, but I can't seem to get Bottle to
    # let me procure a correct JSON response with out using a dictionary.
    # I tried using BaseResponse.  This could be my weakness
    # with Python or confusion in Bottle.
    d = ast.literal_eval(content)
    return d

def callsearch(acsrf,ses_id,psc_pattern,search_string,max_results,remote_addr,targetpage):
    print 'callsearch'
    if targetpage == 'main':
       d = search(acsrf,ses_id,psc_pattern,search_string,max_results,remote_addr,"N")
    else:
       portfolio = request.query['portfolio']
       r = requests.get(URL_TO_MORRIS_PORTFOLIOS_API+"/decoration/"+portfolio)
       content = r.text
       d = ast.literal_eval(r.text)
       p3ids = d['data']
       readCredentials()

       payload = { 'username' : PricesPaidAPIUsername,\
                                        'password' : PricesPaidAPIPassword,\
                                        'p3ids' : pickle.dumps(p3ids)
                        }

       r = requests.post(URLToPPSearchApiSolr+"/fromIds", data=payload, \
                                  auth=(PricesPaidAPIBasicAuthUsername, PricesPaidAPIBasicAuthPassword), verify=False)
       P3Auth.LogActivity.logDebugInfo("Excel Got Past Post to :"+URLToPPSearchApiSolr,request.remote_addr)
       content = r.text
       d = ast.literal_eval(content)
    del d['numFound']
    del d['searchstring']
    return d

@app.wrap_app.route('/SimpleWORD',method='GET')
def createword():
    acsrf = request.query['antiCSRF']
    ses_id = request.query['session_id']
    portfolio = request.query['portfolio']
    targetpage = request.query['targetpage']
    psc_pattern = request.query['psc_pattern']
    psc_pattern = None
    search_string = request.query['search_string']
    max_results = request.query['max_results']

    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf,request.remote_addr)):
        return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                       extra_login_methods=EXTRA_LOGIN_METHODS,
                       footer_html=FOOTER_HTML,
                           goog_anal_script=GoogleAnalyticsInclusionScript)
    d = callsearch(acsrf,ses_id,psc_pattern,search_string,max_results,request.remote_addr,targetpage)
    ts = time.strftime("%d_%m_%y_%H_%M_%S")
    ts1 = time.strftime("%a, %d %b:%H:%M:%S")
    from docx import Document
    document = Document()
    #document.add_heading(u"My title", 0)
    #worksheet.write(row,0,"p3 Search Results",boldformat)
    #worksheet.write(row,0,"This Report is generated by P3 for {0} - on {1} ".format(search_string,ts1),boldformat)
    p = document.add_paragraph('')
    p.add_run('Price Paid Portal Search Results').bold = True
    p = document.add_paragraph('')
    dis = 'This Report is generated by P3 for {0} - on {1} '.format(search_string,ts1)
    p.add_run(dis).bold = True
    #document.add_paragraph('This Report is generated by P3 for {0} - on {1} '.format(search_string,ts1))
    reqd_header_columns = []
    for each, valuesdict in d.iteritems():
       reqd_header_columns.append(['Order Date','orderDate'])
       reqd_header_columns.append(['Product Description','productDescription'])
       reqd_header_columns.append(['Manufacturer Name','Manufacturer Name'])
       reqd_header_columns.append(['Manufacturer Part Number','Manufacturer Part Number'])
       reqd_header_columns.append(['Unit Price','unitPrice'])
       reqd_header_columns.append(['Unit of Issue','Unit of Issue'])
       reqd_header_columns.append(['Qty','Quantity'])
       reqd_header_columns.append(['Extended Price','Extended Price'])
       reqd_header_columns.append(['Ordering Agency','contractingAgency'])
       reqd_header_columns.append(['Ordering Bureau','Bureau'])
       reqd_header_columns.append(['Contract Number','Contract Number'])
       reqd_header_columns.append(['Vendor','vendor'])
       reqd_header_columns.append(['Data Source','dataSource'])
       break
    for key, vdict in d.iteritems():
        for each in vdict:
            if each in ["unitPrice","Extended Price"]:
                vdict[each] = '$' + vdict[each][:-2]

    for key, vdict in d.iteritems():
        table = document.add_table(rows=14,cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field Name'
        hdr_cells[1].text = 'Value'
        for each in range(len(reqd_header_columns)):
                row_cells = table.rows[each+1].cells
 		row_cells[0].text = reqd_header_columns[each][0]
 		row_cells[1].text = vdict[reqd_header_columns[each][1]]
	document.add_paragraph('')
                #try:
                #    worksheet.write(row,each,vdict[reqd_header_columns[each][1]])
                #except:
                #    worksheet.write(row,each,' ')
        #row += 1


    import StringIO
    output = StringIO.StringIO()
    document.save(output)

    filename = "sampeledoc.doc"

    output.seek(0)
    filename = 'WORD_Report_' + portfolio + '_' + ts
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document; charset=UTF-8'
    response.headers['Content-Disposition'] = 'attachment; filename='+filename
    return output.read()



@app.wrap_app.route('/SimpleCSV',method='GET')
def createxls():
    acsrf = request.query['antiCSRF']
    ses_id = request.query['session_id']
    portfolio = request.query['portfolio']
    targetpage = request.query['targetpage']
    psc_pattern = request.query['psc_pattern']
    psc_pattern = None
    search_string = request.query['search_string']
    max_results = request.query['max_results']

    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf,request.remote_addr)):
        return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                       extra_login_methods=EXTRA_LOGIN_METHODS,
                       footer_html=FOOTER_HTML,
                           goog_anal_script=GoogleAnalyticsInclusionScript)
    d = callsearch(acsrf,ses_id,psc_pattern,search_string,max_results,request.remote_addr,targetpage)
    xls = ""
    import StringIO
    output = StringIO.StringIO()

    ts = time.strftime("%d_%m_%y_%H_%M_%S")
    ts1 = time.strftime("%a, %d %b:%H:%M:%S")
    filename = 'XLS_Report_' + portfolio + '_' + ts
    workbook = xlsxwriter.Workbook(output, {'in_memory':True})
    worksheet = workbook.add_worksheet()
    boldformat = workbook.add_format({'bold':True})
    create_hdrs = True
    row = 0
    worksheet.write(row,0,"p3 Search Results",boldformat)
    row += 1
    worksheet.write(row,0,"This Report is generated by P3 for {0} - on {1} ".format(search_string,ts1),boldformat)
    row = 3
    reqd_header_columns = []
    for each, valuesdict in d.iteritems():
       reqd_header_columns.append(['Order Date','orderDate'])
       reqd_header_columns.append(['Product Description','productDescription'])
       reqd_header_columns.append(['Manufacturer Name','Manufacturer Name'])
       reqd_header_columns.append(['Manufacturer Part Number','Manufacturer Part Number'])
       reqd_header_columns.append(['Unit Price','unitPrice'])
       reqd_header_columns.append(['Unit of Issue','Unit of Issue'])
       reqd_header_columns.append(['Qty','Quantity'])
       reqd_header_columns.append(['Extended Price','Extended Price'])
       reqd_header_columns.append(['Ordering Agency','contractingAgency'])
       reqd_header_columns.append(['Ordering Bureau','Bureau'])
       reqd_header_columns.append(['Contract Number','Contract Number'])
       reqd_header_columns.append(['Vendor','vendor'])
       reqd_header_columns.append(['Data Source','dataSource'])
       break

    worksheet.set_column(0,15,25)
    #print 'd.iteritems() =', d.iteritems()
    for key, vdict in d.iteritems():
	for each in vdict:
            if each in ["unitPrice","Extended Price"]:
		vdict[each] = '$' + vdict[each][:-2] 
	    
    for key, vdict in d.iteritems():
        if create_hdrs:
                for each in range(len(reqd_header_columns)):
                        worksheet.write(row,each,reqd_header_columns[each][0],boldformat)
                row += 1
        create_hdrs = False

        for each in range(len(reqd_header_columns)):
                try:
                    worksheet.write(row,each,vdict[reqd_header_columns[each][1]])
 		except:
		    worksheet.write(row,each,' ')
	row += 1
    workbook.close()
    output.seek(0)
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet; charset=UTF-8'
    response.headers['Content-Disposition'] = 'attachment; filename='+filename

    return output.read()


@app.wrap_app.route('/SimpleHTML',method='GET')
def apihtml():
    acsrf = request.query['antiCSRF']
    ses_id = request.query['session_id']
    portfolio = request.query['portfolio']
    targetpage = request.query['targetpage']
    psc_pattern = request.query['psc_pattern']
    psc_pattern = None
    search_string = request.query['search_string']
    max_results = request.query['max_results']
    #print 'psc_pattern =', psc_pattern
    #The value returned from js is undefined, thats not the case for post request, Need to check...For now, it should not hamper the results

    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf,request.remote_addr)):
        return template('Login',message='Improper Credentials or Timeout.',
			 ombmax_redirect_url=OMBMAX_REDIRECT_URL,
                       extra_login_methods=EXTRA_LOGIN_METHODS,
                       footer_html=FOOTER_HTML,
                           goog_anal_script=GoogleAnalyticsInclusionScript)

    d = callsearch(acsrf,ses_id,psc_pattern,search_string,max_results,request.remote_addr,targetpage)

    html = ""
   
    for key, vdict in d.iteritems():
	for each in vdict:
            if each in ["unitPrice","Extended Price"]:
		vdict[each] = '$' + vdict[each][:-2] 
    for key, vdict in d.iteritems():
        html = html + produceHTML(vdict)
    # Actually, here we need to loop over a template, but I will try this first!
    #print 'returned response =', response
    return html


def getValue(valuesdict,item):
    try:
        value = valuesdict[item]
    except:
        value = ""
    return value

def produceHTML(valuesdict):
    html = ""

    html = html + "<h3 style= 'display:inline'>"+"Order Date: "+"</h3>"+ "<p style='display:inline'>"+ getValue(valuesdict,"orderDate")+"</p>"+"<br />"
    html = html + "<h3 style= 'display:inline'>"+"Product Description: "+"</h3>"+ "<p style='display:inline'>"+ getValue(valuesdict,"productDescription")+"</p>"+"<br />"
    html = html + "<h3 style= 'display:inline'>"+"Manufacturer Name: "+"</h3>"+ "<p style='display:inline'>"+ getValue(valuesdict,"Manufacturer Name")+"</p>"+"<br />"
    html = html + "<h3 style= 'display:inline'>"+"Manufacturer Part Number: "+"</h3>"+ "<p style='display:inline'>"+ getValue(valuesdict,"Manufacturer Part Number")+"</p>"+"<br />"
    html = html + "<h3 style= 'display:inline'>"+"Unit Price: "+"</h3>"+ "<p style='display:inline'>" + getValue(valuesdict,"unitPrice")+"</p>"+"<br />"
    html = html + "<h3 style= 'display:inline'>"+"Unit of Issue: "+"</h3>"+ "<p style='display:inline'>" + getValue(valuesdict,"Unit of Issue")+"</p>"+"<br />"
    
    html = html + "<h3 style= 'display:inline'>"+"Qty: "+"</h3>"+ "<p style='display:inline'>" + getValue(valuesdict,"Quantity")+"</p>"+"<br />"

    html = html + "<h3 style= 'display:inline'>"+"Extended Price: "+"</h3>"+ "<p style='display:inline'>"+ getValue(valuesdict,"Extended Price")+"</p>"+"<br />"

    html = html + "<h3 style= 'display:inline'>"+"Ordering Agency: "+"</h3>"+ "<p style='display:inline'>"+ getValue(valuesdict,"contractingAgency")+"</p>"+"<br />"

    html = html + "<h3 style= 'display:inline'>"+"Ordering Bureau: "+"</h3>"+ "<p style='display:inline'>"+ getValue(valuesdict,"Bureau")+"</p>"+"<br />"

    html = html + "<h3 style= 'display:inline'>"+"Contract Number: "+"</h3>"+ "<p style='display:inline'>"+ getValue(valuesdict,"Contract Number")+"</p>"+"<br />"

    html = html + "<h3 style= 'display:inline'>"+"Vendor: "+"</h3>"+ "<p style='display:inline'>"+ getValue(valuesdict,"vendor")+"</p>"+"<br />"

    html = html + "<h3 style= 'display:inline'>"+"Data Source: "+"</h3>"+ "<p style='display:inline'>"+ getValue(valuesdict,"dataSource")+"</p>"+"<br />"
    #for k,v in valuesdict.iteritems():
    #    if k not in ("unitPrice","longDescription" ,"productDescription" , "unitsOrdered" , "vendor", "score", "orderDate", "p3id", 'awardIdIdv', 'psc', "contractingAgency"):
    #        html = html + "<h3 style= 'display:inline'>" +k+ ":" +"</h3>"+ "<p style='display:inline'>"+str(v) + "</p>" +"<br />"+ "\n"
    html = html + "<p></p>" + "\n"
    return html

@app.wrap_app.route('/search',method='POST')
def apisolr():
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')
    search_string = request.forms.get('search_string')
    psc_pattern = request.forms.get('psc_pattern')
    max_results = request.forms.get('numRows')
    main_search = request.forms.get('main_search')
    
    return search(acsrf,ses_id,psc_pattern,search_string,max_results,request.remote_addr,main_search)

def search(acsrf,ses_id,psc_pattern,search_string,max_results,remote_addr,main_search):
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf,remote_addr)):
        dict = {0: {"status": "BadAuthentication"}}
        return dict;
    P3Auth.auth.update_acsrf(ses_id,remote_addr)

    search_string = P3Search.check_search_list(ses_id,search_string,main_search)
    readCredentials()

    P3Auth.LogActivity.logSearchBegun(ses_id,psc_pattern,search_string,remote_addr)

    payload = { 'username' : PricesPaidAPIUsername,\
                                'password' : PricesPaidAPIPassword,\
                                'search_string': search_string,\
                                'psc_pattern': psc_pattern,\
                                'numRows': max_results,\
				'remote_addr':request.remote_addr }
    r = requests.post(URLToPPSearchApiSolr, data=payload, \
                          auth=(PricesPaidAPIBasicAuthUsername, PricesPaidAPIBasicAuthPassword), verify=False)

    P3Auth.LogActivity.logDebugInfo("Got Past Post to :"+URLToPPSearchApiSolr,remote_addr)

    content = r.text
    # This is inefficient, but I can't seem to get Bottle to
 # let me procure a correct JSON response with out using a dictionary.
    # I tried using BaseResponse.  This could be my weakness
    # with Python or confusion in Bottle.
    d = ast.literal_eval(content)
    d["searchstring"] = search_string
    print 'content = ', d
    #print 'results returned =', d
    P3Auth.LogActivity.logSearchDone(ses_id,psc_pattern,search_string,remote_addr)
    return d

@app.wrap_app.route('/record_feedback',method='POST')
def feedback():
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')
    message = request.forms.get('message')
    print "Feedback message = ", message
    print "acsrf = ", acsrf
    print "ses_id = ", ses_id
    name = request.forms.get('name')
    name = request.forms.get('name')
    radio_list_value = request.forms.get('radio_list_value')
   
    P3Auth.LogActivity.logDebugInfo("acsrf ses_d :"+acsrf+ses_id,request.remote_addr)
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf,request.remote_addr)):
        dict = {0: {"status": "BadAuthentication"}}
	print 'authentication failed'
        return dict;
    P3Auth.LogActivity.logDebugInfo("authenticated !",request.remote_addr)
    P3Auth.LogActivity.logFeedback(ses_id,request.remote_addr)
    LogFeedback.logFeedback(name,message,radio_list_value,request.remote_addr);
    return "true";


@error(405)
def error405(error):
    return 'The pages are all secured'
# This file is a directoy copy fromt he MorrisData Decorator.  It 
# out to be possible to avoid this duplication, but I don't really 
# know how to do that in Python.  I will have to spend the next day 
# in refactoring.

# BEGIN SERVER-SIDE CALLS TO MORRIS PORTFOLIO API
@app.wrap_app.route('/portfolio', method='GET')
def get_portfolios():
    P3Auth.LogActivity.logDebugInfo("Begin Get Portfolios")
    acsrf = request.query['antiCSRF']
    ses_id = request.query['session_id']
    P3Auth.LogActivity.logDebugInfo("Gotten on Portfolio: acsrf ses_id :"+acsrf+","+ses_id)
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        dict = {0: {"status": "BadAuthentication"}}
        P3Auth.LogActivity.logDebugInfo(" BadAuthentication :"+acsrf+","+ses_id)
        return dict;
    print "Prem URL_TO_MORRIS_PORTFOLIOS_API = ",  URL_TO_MORRIS_PORTFOLIOS_API
    r = requests.get(URL_TO_MORRIS_PORTFOLIOS_API+"/decoration")
    d = ast.literal_eval(r.text)
    return d

@app.wrap_app.route('/portfolio/<name>', method='GET')
def get_specific_tags(name):
    print "Prem Adding PortFolios"
    acsrf = request.query['antiCSRF']
    ses_id = request.query['session_id']
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        dict = {0: {"status": "BadAuthentication"}}
        return dict;
    print "Prem URL_TO_MORRIS_PORTFOLIOS_API = ",  URL_TO_MORRIS_PORTFOLIOS_API
    r = requests.get(URL_TO_MORRIS_PORTFOLIOS_API+"/content/"+name)
    return r.text

@app.wrap_app.route('/portfolio/<name>', method='POST')
def get_create_portfolio(name):
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')

    P3Auth.LogActivity.logDebugInfo("acsrf ses_d :"+acsrf+ses_id)
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        dict = {0: {"status": "BadAuthentication"}}
        return dict;
    #Validate if the portfolio name contains characters other than alphanumeric
    m_obj = re.search(r"^[a-zA-Z0-9]+$",name)
    if (len(name) <= 10 and m_obj is not None):
	r = requests.post(URL_TO_MORRIS_PORTFOLIOS_API+"/decoration/"+name)
    else:
	print "Invalid name in portfolio characters = ", name
        return None
    return r.text

@app.wrap_app.route('/portfolio_export', method='GET')
def get_export_portfolio():
    acsrf = request.query['antiCSRF']
    ses_id = request.query['session_id']
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        dict = {0: {"status": "BadAuthentication"}}
        return dict;
    r = requests.get(URL_TO_MORRIS_PORTFOLIOS_API+"/decoration_export")
    return r.text

@app.wrap_app.route('/portfolio_records', method='GET')
def get_records():
    acsrf = request.query['antiCSRF']
    ses_id = request.query['session_id']

    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        dict = {0: {"status": "BadAuthentication"}}
        return dict;
    r = requests.get(URL_TO_MORRIS_PORTFOLIOS_API+"/decoration_records")
    return r.text

@app.wrap_app.route('/portfolio_records_with_cd/<columns>', method='GET')
def get_records(columns):
    acsrf = request.query['antiCSRF']
    ses_id = request.query['session_id']

    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        dict = {0: {"status": "BadAuthentication"}}
        return dict;
    r = requests.get(URL_TO_MORRIS_PORTFOLIOS_API+"/content_records_with_client_data/"+columns)
    return r.text

@app.wrap_app.route('/portfolio/add_record/<portfolio>/<key>',method='POST')
def add_record_to_portfolio(key,portfolio):
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')

    P3Auth.LogActivity.logDebugInfo("acsrf ses_d :"+acsrf+ses_id)
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        dict = {0: {"status": "BadAuthentication"}}
        return dict;
    r = requests.post(URL_TO_MORRIS_PORTFOLIOS_API+"/decoration/add_record/"+portfolio+"/"+key)
    return r.text

@app.wrap_app.route('/portfolio/delete_decoration/<portfolio>',method='POST')
def delete_portfolio(portfolio):
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')

    P3Auth.LogActivity.logDebugInfo("acsrf ses_d :"+acsrf+ses_id)
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        dict = {0: {"status": "BadAuthentication"}}
        return dict;
    r = requests.post(URL_TO_MORRIS_PORTFOLIOS_API+"/delete_decoration/"+portfolio)
    return r.text

@app.wrap_app.route('/portfolio/delete_association/<portfolio>/<transaction>',method='POST')
def delete_association(portfolio,transaction):
    acsrf = request.forms.get('antiCSRF')
    ses_id = request.forms.get('session_id')

    P3Auth.LogActivity.logDebugInfo("acsrf ses_d :"+repr(acsrf)+' '+repr(ses_id))
    if (not P3Auth.auth.is_valid_acsrf(ses_id,acsrf)):
        dict = {0: {"status": "BadAuthentication"}}
        return dict;
    r = requests.post(URL_TO_MORRIS_PORTFOLIOS_API+"/delete_association/"+portfolio+"/"+transaction)
    return r.text

@app.wrap_app.route('/bubble',method='GET')
def bubble():
    return template('bubble')   

def bytesto(bytes, to, bsize=1024):
    """convert bytes to megabytes, etc.
       sample code:
           print('mb= ' + str(bytesto(314575262000000, 'm')))

       sample output: 
           mb= 300002347.946
    """

    a = {'k' : 1, 'm': 2, 'g' : 3, 't' : 4, 'p' : 5, 'e' : 6 }
    r = float(bytes)
    for i in range(a[to]):
        r = r / bsize

    return(r)
# End Portfolio work

